from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def build_resume_docx(optimized_resume_text: str, output_path: str):
    """
    Converts plain optimized resume text into a
    professionally formatted .docx file.
    """
    doc = Document()

    # ── Page margins ──────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # ── Split text into lines ──────────────────
    lines = optimized_resume_text.strip().split("\n")

    for i, line in enumerate(lines):
        line = line.strip()

        # Skip empty lines — add spacing instead
        if not line:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # ── Detect Name (first non-empty line) ──
        if i == 0 or (i < 3 and not any(
            kw in line.lower() for kw in
            ["email", "phone", "github", "linkedin", "@", "|"]
        ) and len(line.split()) <= 5):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name  = "Calibri"
            run.font.size  = Pt(20)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # ── Detect Contact Info line ──
        elif any(kw in line.lower() for kw in
                 ["@", "github", "linkedin", "phone", "|", "http"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name  = "Calibri"
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        # ── Detect Section Headers ──
        elif is_section_header(line):
            # Add a line above section
            p = doc.add_paragraph()
            add_horizontal_line(p)

            p2 = doc.add_paragraph()
            run = p2.add_run(line.upper().replace("**", "").replace(":", ""))
            run.font.name  = "Calibri"
            run.font.size  = Pt(12)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            p2.paragraph_format.space_before = Pt(6)
            p2.paragraph_format.space_after  = Pt(4)

        # ── Detect Bullet Points ──
        elif line.startswith(("-", "•", "*", "+")):
            clean = line.lstrip("-•*+ ").strip()
            clean = clean.replace("**", "")
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_after  = Pt(2)

        # ── Detect Sub-bullets (STAR format) ──
        elif line.startswith(("\t-", "\t•", "  -", "  •")):
            clean = line.strip().lstrip("-•*+ ").strip()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(1)

        # ── Bold lines (project titles, job titles) ──
        elif line.startswith("**") and line.endswith("**"):
            clean = line.replace("**", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            p.paragraph_format.space_before = Pt(4)

        # ── Regular paragraph ──
        else:
            clean = line.replace("**", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            p.paragraph_format.space_after = Pt(2)

    doc.save(output_path)
    return output_path


def is_section_header(line: str) -> bool:
    """Detects if a line is a resume section header."""
    headers = [
        "summary", "objective", "education", "experience",
        "skills", "technical skills", "projects", "certifications",
        "achievements", "awards", "languages", "interests",
        "work experience", "professional experience"
    ]
    clean = line.lower().replace("**", "").replace(":", "").strip()
    return any(clean == h or clean == h.upper() for h in headers)


def add_horizontal_line(paragraph):
    """Adds a thin horizontal line above section headers."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)